# %% [markdown]
# # Build journal-formatted DOCX files
#
# The script converts the verified Markdown manuscript and appendix into simple,
# line-numbered Word documents for initial journal submission.

# %%
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def enable_line_numbers(section):
    sect_pr = section._sectPr
    line_num = OxmlElement("w:lnNumType")
    line_num.set(qn("w:countBy"), "1")
    line_num.set(qn("w:restart"), "newPage")
    sect_pr.append(line_num)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def add_markdown(document: Document, markdown: str, include_title: bool = True) -> None:
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        text = clean_inline(" ".join(part.strip() for part in buffer))
        buffer.clear()
        if text:
            paragraph = document.add_paragraph(text, style="Body Text")
            paragraph.paragraph_format.keep_together = False

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1 and include_title:
                paragraph = document.add_paragraph(style="Title")
                paragraph.add_run(text)
            elif level == 1:
                paragraph = document.add_paragraph(text, style="Heading 1")
            else:
                paragraph = document.add_paragraph(text, style=f"Heading {min(level - 1, 3)}")
            continue
        if line.startswith("- "):
            flush()
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            flush()
            document.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            continue
        buffer.append(line)
    flush()


def configure(document: Document, short_title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    enable_line_numbers(section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    body = styles["Body Text"]
    body.font.name = "Arial"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.first_line_indent = Inches(0)

    for style_name, size in (("Title", 15), ("Heading 1", 13), ("Heading 2", 12), ("Heading 3", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    header = section.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    add_page_number(section.footer.paragraphs[0])

    core = document.core_properties
    core.title = short_title
    core.subject = "Submission manuscript"
    core.author = "[AUTHOR INPUT NEEDED]"
    core.keywords = "clinical prediction; calibration; outcome ascertainment; acute kidney injury"
    core.comments = "Statistical values require author verification before submission."


def build_main(manuscript: Path, figure_legends: Path, table_legends: Path, output: Path) -> None:
    document = Document()
    configure(document, "Outcome-measurement transport in clinical AI")
    add_markdown(document, manuscript.read_text(encoding="utf-8"), include_title=True)
    document.add_page_break()
    add_markdown(document, figure_legends.read_text(encoding="utf-8"), include_title=False)
    document.add_page_break()
    add_markdown(document, table_legends.read_text(encoding="utf-8"), include_title=False)
    document.save(output)


def build_supplement(appendix: Path, output: Path) -> None:
    document = Document()
    configure(document, "Supplementary appendix: outcome-measurement transport")
    add_markdown(document, appendix.read_text(encoding="utf-8"), include_title=True)
    document.save(output)


def build_cover_letter(letter: Path, output: Path) -> None:
    document = Document()
    configure(document, "Cover letter: outcome-measurement transport")
    section = document.sections[0]
    # Cover letters do not need manuscript line numbering or a running page footer.
    line_numbering = section._sectPr.find(qn("w:lnNumType"))
    if line_numbering is not None:
        section._sectPr.remove(line_numbering)
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""
    document.styles["Body Text"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    document.styles["Body Text"].paragraph_format.space_after = Pt(8)
    add_markdown(document, letter.read_text(encoding="utf-8"), include_title=False)
    document.save(output)


if __name__ == "__main__":
    source = Path(sys.argv[1])
    out = Path(sys.argv[2])
    build_main(
        source / "MANUSCRIPT_FINAL.md",
        source / "FIGURE_LEGENDS.md",
        source / "TABLE_LEGENDS.md",
        out / "MANUSCRIPT_LANCET_DIGITAL_HEALTH.docx",
    )
    build_supplement(
        source / "SUPPLEMENTARY_APPENDIX.md",
        out / "SUPPLEMENTARY_APPENDIX.docx",
    )
    build_cover_letter(
        source / "COVER_LETTER_LANCET_DIGITAL_HEALTH.md",
        out / "COVER_LETTER_LANCET_DIGITAL_HEALTH.docx",
    )
